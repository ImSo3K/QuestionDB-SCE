import examSystemData
import textData
import os
import glob
from docx import Document
from docx.shared import Cm
import tempfile
from pdf2image import convert_from_path, convert_from_bytes
from pdf2image.exceptions import (
    PDFInfoNotInstalledError,
    PDFPageCountError,
    PDFSyntaxError
)


def init_and_menu():
    obj1 = examSystemData.Exam()
    textData.get_exam_info(obj1)
    textData.get_question_info(obj1)
    obj2 = examSystemData.Course()
    textData.get_course_info(obj2, obj1)
    obj3 = examSystemData.Department()
    textData.get_department_info(obj2, obj3)
    print('Press 1 to print Department data')
    print('Press 2 to print Course data')
    print('Press 3 to print Exam data')
    print('Press 4 to convert pdf2image & then image2docx')
    print('Press 0 to exit')
    choice = 1
    while 1 <= choice <= 4:
        choice = int(input())

        if choice == 1:
            print(f'Department Name : {obj3.dep_name}\n' + f'Department Coordinator Name : {obj3.coordinator_name}')

        elif choice == 2:
            print(f'Course Name : {obj2.course_name}\n' + f'Course Lecturer Name : {obj2.lecturer_name}')
            print(f'Year Of Course : {obj2.year_of_course}')

        elif choice == 3:
            print(f'Exam Date : {obj2.exam.exam_date}\n' + f'Exam semester : {obj2.exam.exam_semester}')
            print('Question list :\n')
            for i in obj2.exam.question_list:
                examSystemData.Question.print_question_info(i)

        elif choice == 4:
            pdf_2_image()
            image_2_docx()

        elif choice == 0:
            break

        else:
            print('Invalid Choice')


def pdf_2_image():
    with tempfile.TemporaryDirectory() as path:
        images_from_path = convert_from_path('/home/sharon/Desktop/Python/PyCharm/Project1/pdf_file/BW.pdf',
                                             output_folder='/home/sharon/Desktop/Python/PyCharm/Project1'
                                                           '/images_from_pdf_demo')


def image_2_docx():
    files_name = []
    os.chdir("/home/sharon/Desktop/Python/PyCharm/Project1/images_from_pdf_demo")
    for file in glob.glob("*.png"):
        files_name.append(file)
    files_name.sort()

    document = Document()

    for i in range(0, len(files_name), 2):
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(f'/home/sharon/Desktop/Python/PyCharm/Project1/images_from_pdf_demo/{files_name[i]}',
                      width=Cm(15), height=Cm(17))

    document.save('/home/sharon/Desktop/Python/PyCharm/Project1/docx_demo/demo.docx')


print('initializing objects with the data from the given text files')
init_and_menu()
