import os
import glob
from docx import Document
from docx.shared import Cm


def image_2_docx():
    files_name = []
    os.chdir("/home/sharon/Desktop/Python/PyCharm/Project1/images_from_pdf_demo")
    for file in glob.glob("*.png"):
        files_name.append(file)
    files_name.sort()

    document = Document()

    for i in range(0, len(files_name), 2):
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(f'/home/sharon/Desktop/Python/PyCharm/Project1/images_from_pdf_demo/{files_name[i]}',
                      width=Cm(15), height=Cm(17))

    document.save('/home/sharon/Desktop/Python/PyCharm/Project1/docx_demo/demo.docx')
