from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_text('Good Morning every body,This is my ')
r.add_picture('/tmp/foo.jpg')
r.add_text(' do you like it?')

document.save('demo.docx')